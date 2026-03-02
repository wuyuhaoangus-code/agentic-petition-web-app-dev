import io
import re
import logging
import subprocess
import tempfile
import os
from typing import AsyncGenerator, Dict, Any

from google.cloud import vision
from pdf2image import convert_from_bytes
import docx
import filetype

logger = logging.getLogger(__name__)

class TextProcessor:
    """
    Handles ONLY the cleaning and formatting of text (Single Responsibility Principle).
    This logic is now decoupled from the OCR provider.
    """
    @staticmethod
    def clean(text: str) -> str:
        if not text:
            return ""

        # 1. Normalize Whitespace (Horizontal)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # 2. Remove "OCR Noise" (Lone symbols from creases/dust)
        text = re.sub(r'(?m)^[·•|~.:_-]$', '', text)

        # 3. Structural Reconstruction
        lines = text.splitlines()
        cleaned_lines = []
        
        for line in lines:
            stripped = line.strip()
            if stripped:
                cleaned_lines.append(stripped)
            # Keep exactly ONE empty line between blocks for LLM paragraph detection
            elif cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")

        # 4. Final Polish (Limit max vertical gaps)
        content = "\n".join(cleaned_lines).strip()
        return re.sub(r'\n{3,}', '\n\n', content)


class VisionService:
    """
    Handles the orchestration of PDF, Image, and Word processing.
    Follows Dependency Inversion by accepting a client and processor.
    """
    def __init__(self, client: vision.ImageAnnotatorClient = None, processor: TextProcessor = None):
        self._client = client
        self.processor = processor or TextProcessor()

    @property
    def client(self) -> vision.ImageAnnotatorClient:
        """Lazy initialization of the Google Vision client."""
        if self._client is None:
            try:
                self._client = vision.ImageAnnotatorClient()
            except Exception as e:
                logger.error(f"Failed to initialize Google Vision client: {e}")
                raise
        return self._client

    def _image_to_bytes(self, image) -> bytes:
        """Helper to convert PIL image to JPEG bytes."""
        img_byte_arr = io.BytesIO()
        image.save(img_byte_arr, format='JPEG')
        return img_byte_arr.getvalue()

    async def extract_text_stream(self, file_content: bytes, max_pages: int = 20) -> AsyncGenerator[Dict[str, Any], None]:
        """
        Detects file type and generates cleaned text.
        """
        kind = filetype.guess(file_content)
        mime = kind.mime if kind else "application/octet-stream"
        
        logger.info(f"--- Processing New File ---")
        logger.info(f"Detected MIME type: {mime}")

        if mime == "application/pdf":
            logger.info("Starting PDF processing...")
            async for result in self._process_pdf(file_content, max_pages):
                yield result
        elif mime.startswith("image/"):
            logger.info(f"Starting image processing for {mime}...")
            async for result in self._process_image(file_content):
                yield result
        elif mime in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/zip"]:
            logger.info("Starting Word (.docx) processing...")
            async for result in self._process_docx(file_content):
                yield result
        elif mime == "application/msword":
            logger.info("Starting Legacy Word (.doc) processing via antiword...")
            async for result in self._process_doc(file_content):
                yield result
        else:
            logger.warning(f"Unsupported file type: {mime}")
            yield {"error": f"Unsupported file format: {mime}"}
        
        logger.info(f"--- File Processing Complete ---")

    async def _process_pdf(self, content: bytes, max_pages: int) -> AsyncGenerator[Dict[str, Any], None]:
        try:
            images = convert_from_bytes(content, dpi=200)
        except Exception as e:
            logger.error(f"Failed to convert PDF to images: {e}")
            yield {"error": "Invalid PDF format"}
            return

        total_pages = min(len(images), max_pages)
        for i, image in enumerate(images[:total_pages]):
            img_bytes = self._image_to_bytes(image)
            async for result in self._ocr_image_bytes(img_bytes, i + 1, total_pages):
                yield result

    async def _process_image(self, content: bytes) -> AsyncGenerator[Dict[str, Any], None]:
        """Process a single JPEG/PNG image."""
        async for result in self._ocr_image_bytes(content, 1, 1):
            yield result

    async def _ocr_image_bytes(self, img_bytes: bytes, current_page: int, total_pages: int) -> AsyncGenerator[Dict[str, Any], None]:
        """Actual OCR call to Google Vision."""
        try:
            vision_image = vision.Image(content=img_bytes)
            response = self.client.document_text_detection(image=vision_image)
            
            if response.error.message:
                logger.error(f"Google Vision Error: {response.error.message}")
                yield {"page": current_page, "error": response.error.message}
                return

            raw_text = response.full_text_annotation.text if response.full_text_annotation else ""
            cleaned_page = self.processor.clean(raw_text)

            yield {
                "page": current_page,
                "total": total_pages,
                "text": f"--- Page {current_page} ---\n" + cleaned_page if total_pages > 1 else cleaned_page
            }
        except Exception as e:
            logger.error(f"OCR Error on page {current_page}: {e}")
            yield {"page": current_page, "error": str(e)}

    async def _process_docx(self, content: bytes) -> AsyncGenerator[Dict[str, Any], None]:
        """Extract text from modern Word (.docx) files."""
        try:
            doc = docx.Document(io.BytesIO(content))
            full_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Also handle tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            cleaned_text = self.processor.clean("\n".join(full_text))
            
            yield {
                "page": 1,
                "total": 1,
                "text": cleaned_text
            }
        except Exception as e:
            logger.error(f"Failed to process Word document: {e}")
            yield {"error": "Invalid or corrupted Word document"}

    async def _process_doc(self, content: bytes) -> AsyncGenerator[Dict[str, Any], None]:
        """Extract text from legacy Word (.doc) files using antiword."""
        temp_file_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
                tmp.write(content)
                temp_file_path = tmp.name

            # Run antiword to extract text
            result = subprocess.run(
                ["antiword", temp_file_path],
                capture_output=True,
                text=True,
                check=True
            )
            
            extracted_text = result.stdout
            cleaned_text = self.processor.clean(extracted_text)
            
            yield {
                "page": 1,
                "total": 1,
                "text": cleaned_text
            }
        except subprocess.CalledProcessError as e:
            logger.error(f"Antiword failed to process document: {e.stderr}")
            yield {"error": "Failed to extract text from legacy Word (.doc) file"}
        except Exception as e:
            logger.error(f"Error processing .doc file: {e}")
            yield {"error": f"Error processing .doc file: {str(e)}"}
        finally:
            if temp_file_path and os.path.exists(temp_file_path):
                os.remove(temp_file_path)

# Singleton instance
vision_service = VisionService()
