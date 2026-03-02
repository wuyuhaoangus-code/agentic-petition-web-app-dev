"""
Cover letter document builder for the petition submission.
Structure: date -> address -> Regarding -> listing of forms/files/exhibits -> list of all exhibits.
Leverages document_builder for styling and _get_item_name.
"""
import logging
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Optional
from sqlalchemy.orm import Session

from app.features.petitions.repositories import ExhibitRepository
from app.features.documents.document_builder import document_builder

logger = logging.getLogger(__name__)

# Default USCIS address (I-140 / EB-1A)
DEFAULT_USCIS_ADDRESS_LINES = [
    "USCIS",
    "Attn: I-140",
    "P.O. Box 660128",
    "Dallas, TX 75266-0128",
]

# Standard forms to list before petition letter and exhibits (customize per visa type)
def _default_forms(petitioner_name: str) -> List[str]:
    return [
        "I-140 Filing Fee check $715 and Asylum Program Fee check $300;",
        "Form G-1145, E-Notification of Application/Petition Acceptance;",
        "Form G-28, Notice of Entry of Appearance as Attorney or Accredited Representative;",
        "Form I-140, Immigrant Petition for Alien Workers;",
        f"Petition Letter for {petitioner_name};",
        "The exhibit documents attached in support of this petition:",
    ]


def _format_date(d: date) -> str:
    """e.g. December 31st, 2025"""
    day = d.day
    if 4 <= day <= 20 or 24 <= day <= 30:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return d.strftime(f"%B {day}{suffix}, %Y")


def build_cover_letter(
    db: Session,
    run_id,
    user_id,
    petitioner_name: str,
    application_id=None,
    letter_date: Optional[date] = None,
    uscis_address_lines: Optional[List[str]] = None,
    re_subject: str = "EB-1A Immigrant Petition for Alien of Extraordinary Ability",
    forms_list: Optional[List[str]] = None,
) -> Document:
    """
    Builds a cover letter .docx with:
    date -> address -> Regarding (Re: + Petitioner) -> Dear USCIS Officer ->
    intro sentence -> numbered forms/exhibits list -> EXHIBIT 1 — Title, items; EXHIBIT 2 — ...
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    letter_date = letter_date or date.today()
    address_lines = uscis_address_lines or DEFAULT_USCIS_ADDRESS_LINES
    rest_forms = forms_list if forms_list is not None else _default_forms(petitioner_name)
    # Numbered list: 1. Cover Letter, then 2.-7. (matches ZIP order)
    forms = ["Cover Letter"] + list(rest_forms)

    # 1. Date (centered)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(_format_date(letter_date))
    document_builder._set_font_and_size(run)

    # 2. USCIS address (left, one line per paragraph)
    for line in address_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        document_builder._set_font_and_size(run)
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

    # 3. Re: subject and Petitioner (2 indents from left, no spacing between the two lines)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Re:      ")
    run.bold = True
    document_builder._set_font_and_size(run)
    run2 = p.add_run(re_subject)
    run2.bold = True
    document_builder._set_font_and_size(run2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Petitioner:      ")
    run.bold = True
    document_builder._set_font_and_size(run)
    run2 = p.add_run(petitioner_name)
    run2.bold = True
    document_builder._set_font_and_size(run2)

    # 4. Salutation
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Dear USCIS Officer:")
    document_builder._set_font_and_size(run)

    # 5. Intro sentence
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Please find the following forms and supporting documents for your review:")
    document_builder._set_font_and_size(run)

    # 6. Numbered list (1. Cover Letter, 2.-7. forms/exhibits intro; matches ZIP order)
    for i, form_text in enumerate(forms, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        run_prefix = p.add_run(f"{i}. ")
        run_prefix.bold = True
        document_builder._set_font_and_size(run_prefix)
        run = p.add_run(form_text)
        document_builder._set_font_and_size(run)
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

    # 7. Exhibits for this run (EXHIBIT N — Title, then Copy of ... for each item)
    exhibit_repo = ExhibitRepository(db)
    exhibits = exhibit_repo.get_by_run(user_id, run_id)

    for exhibit in exhibits:
        # Exhibit header: EXHIBIT N — Title
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(6)
        title_p.paragraph_format.space_after = Pt(0)
        clean_title = exhibit.title.strip().strip('"').strip()
        header_text = f"EXHIBIT {exhibit.exhibit_number} — {clean_title}"
        run_title = title_p.add_run(header_text)
        run_title.bold = True
        document_builder._set_font_and_size(run_title)

        items = exhibit_repo.get_items_for_exhibit_ordered(exhibit.id)
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(0)
            run_prefix = p.add_run(f"{item.item_suffix}. ")
            run_prefix.bold = True
            document_builder._set_font_and_size(run_prefix)
            name = document_builder._get_item_name(db, item)
            run_desc = p.add_run(f"Copy of {name}")
            document_builder._set_font_and_size(run_desc)

    # Optional closing: EB-1A summary paragraph, then thank-you and sign-off
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "By a preponderance of the evidence, the above documents establish that "
        f"{petitioner_name} qualifies for classification as an alien of extraordinary ability "
        "under 8 C.F.R. § 204.5(h)(2) and satisfies at least three of the ten criteria set forth "
        "at 8 C.F.R. § 204.5(h)(3), and that the petitioner’s entry will substantially benefit "
        "prospectively the United States."
    )
    document_builder._set_font_and_size(run)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "Thank you for your time and attention to this matter. "
        "Should you have any further questions on this petition, please do not hesitate to contact us."
    )
    document_builder._set_font_and_size(run)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("Very truly yours,")
    document_builder._set_font_and_size(run)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(petitioner_name)
    document_builder._set_font_and_size(run)

    return doc


def build_cover_letter_bytes(db: Session, run_id, user_id, petitioner_name: str, **kwargs) -> bytes:
    """Build cover letter document and return as bytes for upload."""
    import io

    doc = build_cover_letter(db, run_id=run_id, user_id=user_id, petitioner_name=petitioner_name, **kwargs)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
