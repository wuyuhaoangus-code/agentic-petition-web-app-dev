import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re
from urllib.parse import unquote, urlparse, parse_qs
from typing import List
from sqlalchemy.orm import Session
from app.features.documents.models import UserExhibit, UserExhibitItem, Citation
from app.features.documents.repositories import DocumentRepository, CitationRepository
from app.features.petitions.repositories import ExhibitRepository
from app.features.documents.legal_authority_template import LEGAL_AUTHORITY_PARAGRAPHS

logger = logging.getLogger(__name__)

# Official EB-1A Regulatory Language Mapping for USCIS
CRITERIA_INFO = {
    "awards": {
        "header": "AWARDS",
        "reg_text": "Evidence of the Petitioner’s receipt of lesser nationally or internationally recognized prizes or awards for excellence in the field of endeavor. 8 C.F.R. § 204.5(h)(3)(i)."
    },
    "membership": {
        "header": "MEMBERSHIPS",
        "reg_text": "Evidence of the Petitioner’s membership in associations in the field for which classification is sought, which require outstanding achievements of their members, as judged by recognized national or international experts in their disciplines or fields. 8 C.F.R. § 204.5(h)(3)(ii)."
    },
    "press": {
        "header": "PRESS",
        "reg_text": "Published material about the Petitioner in professional or major trade publications or other major media, relating to the Petitioner’s work in the field for which classification is sought. Such evidence shall include the title, date, and author of the material, and any necessary translation. 8 C.F.R. § 204.5(h)(3)(iii)."
    },
    "published_material": {
        "header": "PUBLISHED MATERIAL",
        "reg_text": "Published material about the Petitioner in professional or major trade publications or other major media, relating to the Petitioner’s work in the field for which classification is sought. Such evidence shall include the title, date, and author of the material, and any necessary translation. 8 C.F.R. § 204.5(h)(3)(iii)."
    },
    "judging": {
        "header": "JUDGING",
        "reg_text": "Evidence of the Petitioner’s participation, either individually or on a panel, as a judge of the work of others in the same or an allied field of specification for which classification is sought. 8 C.F.R. § 204.5(h)(3)(iv)."

    },
    "contributions": {
        "header": "ORIGINAL CONTRIBUTION",
        "reg_text": "Evidence of the Petitioner’s original scientific, scholarly, artistic, athletic, or business-related contributions of major significance in the field. 8 C.F.R. § 204.5(h)(3)(v)."
    },
    "scholarly": {
        "header": "SCHOLARLY ARTICLES",
        "reg_text": "Evidence of the Petitioner’s authorship of scholarly articles in the field, in professional or major trade publications or other major media. 8 C.F.R. § 204.5(h)(3)(vi)."
    },
    "leading": {
        "header": "LEADING ROLE",
        "reg_text": "Evidence that the Petitioner has performed in a leading or critical role for organizations or establishments that have a distinguished reputation. 8 C.F.R. § 204.5(h)(3)(viii)."
    },
    "exhibitions": {
        "header": "ARTISTIC EXHIBITIONS",
        "reg_text": "Evidence of the display of the Petitioner’s work in the field at artistic exhibitions or showcases. 8 C.F.R. § 204.5(h)(3)(vii)."
    },
    "salary": {
        "header": "HIGH SALARY",
        "reg_text": "Evidence that the Petitioner has commanded a high salary or other significantly high remuneration for services, in relation to others in the field. 8 C.F.R. § 204.5(h)(3)(ix)."
    },
    "commercial": {
        "header": "COMMERCIAL SUCCESS",
        "reg_text": "Evidence of commercial successes in the performing arts, as shown by box office receipts or record, cassette, compact disk, or video sales. 8 C.F.R. § 204.5(h)(3)(x)."
    },
    "recommendation": {
        "header": "RECOMMENDATION LETTERS",
        "reg_text": "Recommendation Letters from Industry Experts"
    },
    "future_work": {
        "header": "FUTURE WORK IN THE UNITED STATES",
        "reg_text": "Evidence that the Petitioner will continue to work in the area of expertise in the United States and that their work will substantially benefit prospectively the United States."
    },
}

class DocumentBuilder:
    def _set_font_and_size(self, run, font_name="Times New Roman", font_size=Pt(12)):
        run.font.name = font_name
        run.font.size = font_size
        rPr = run._r.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), font_name)

    def _add_hyperlink(self, paragraph, text, url, color="0000EE", underline=True):
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if color:
            c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
        if underline:
            u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
        new_run.append(rPr)
        text_element = OxmlElement('w:t'); text_element.text = text; new_run.append(text_element)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def _get_item_name(self, db: Session, item: UserExhibitItem) -> str:
        name = "Document"
        repo = DocumentRepository(db)
        if item.content_id:
            content_rec = repo.get_evidence_by_id(item.content_id)
            if content_rec and content_rec.title:
                name = content_rec.title
        elif item.file_id:
            content_rec = repo.get_evidence_for_file_any(item.file_id)
            file_rec = repo.get_file_by_id(item.file_id)
            if content_rec and content_rec.title and len(content_rec.title.strip()) > 2:
                name = content_rec.title
            elif file_rec:
                name = file_rec.file_name
        
        # Clean up automated prefixes but preserve core characters (including Chinese)
        name = re.sub(r'^[a-z0-9]\.\s+', '', str(name), flags=re.IGNORECASE)
        name = name.replace("Copy of ", "").strip()
        return name

    def generate_petition_intro(self, intro_text: str, db: Session = None, user_id: any = None) -> Document:
        """
        Generates the pre-exhibits petition intro (Introduction, Legal Authority, Argument).
        Matches the format of the sample EB1A intro document.
        """
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        if not intro_text:
            return doc

        # Safely convert intro_text to string if it's not already
        if isinstance(intro_text, dict):
            intro_text = str(intro_text.get("text", intro_text.get("content", str(intro_text))))
        elif not isinstance(intro_text, str):
            intro_text = str(intro_text) if intro_text is not None else ""

        # Final safety check - ensure it's a string before calling string methods
        if not isinstance(intro_text, str):
            logger.warning(f"intro_text is still not a string after conversion: {type(intro_text)}, converting to string")
            intro_text = str(intro_text) if intro_text is not None else ""

        if not intro_text:
            return doc

        # Split by one or more newlines that have only whitespace between them
        if isinstance(intro_text, str):
            paragraphs = re.split(r'\n\s*\n', intro_text.strip())
        else:
            logger.error(f"Cannot split intro_text - not a string: {type(intro_text)}")
            paragraphs = [str(intro_text)] if intro_text else []
        
        # Track which section we're in
        in_legal_authority = False
        in_legal_criteria_list = False
        in_argument = False
        
        for para_text in paragraphs:
            # Ensure para_text is a string before calling string methods
            if not isinstance(para_text, str):
                para_text = str(para_text) if para_text is not None else ""
            
            if not para_text.strip():
                continue
            
            # Cleanup and Normalization - remove markdown, fix formatting issues
            clean_para = para_text.strip().replace("**", "").replace("*", "")
            # Remove dict-like prefixes that might appear (more aggressive cleaning)
            clean_para = re.sub(r'^\{\'', '', clean_para)
            clean_para = re.sub(r'^\'', '', clean_para)
            clean_para = re.sub(r'^\{\"', '', clean_para)
            clean_para = re.sub(r'^\"', '', clean_para)
            # Remove any remaining dict artifacts
            clean_para = re.sub(r'^\{', '', clean_para)
            clean_para = clean_para.strip()
            
            if not clean_para:
                continue

            # --- HEADER SECTION FORMATTING ---
            
            # 1. Date (e.g. "December 4th, 2024")
            if re.match(r'^[A-Z][a-z]+\s+\d{1,2}(?:st|nd|rd|th),\s+\d{4}$', clean_para):
                 p = doc.add_paragraph()
                 p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                 p.paragraph_format.space_after = Pt(24)
                 run = p.add_run(clean_para)
                 self._set_font_and_size(run)
                 continue

            # 2. USCIS Address Block
            if (clean_para.startswith("USCIS") or 
                clean_para.startswith("Attn: I-140") or 
                clean_para.startswith("P.O. Box") or 
                re.match(r'^[A-Z][a-zA-Z\s]+,\s+[A-Z]{2}\s+\d{5}', clean_para)):
                 p = doc.add_paragraph()
                 p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                 p.paragraph_format.space_after = Pt(0)
                 run = p.add_run(clean_para)
                 self._set_font_and_size(run)
                 continue

            # 3. Re: Section
            if clean_para.startswith("Re:"):
                 lines = clean_para.split('\n')
                 for idx, line in enumerate(lines):
                     if not line.strip(): continue
                     p = doc.add_paragraph()
                     # Center alignment for Re block matches the style of being distinct
                     p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                     p.paragraph_format.space_before = Pt(12) if idx == 0 else Pt(0)
                     p.paragraph_format.space_after = Pt(0)
                     
                     run = p.add_run(line.strip())
                     run.bold = True
                     self._set_font_and_size(run)
                 
                 # Add space after the Re block
                 if doc.paragraphs:
                     doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
                 continue

            # 4. Salutation
            if clean_para.startswith("Dear Sir/ Madam") or clean_para.startswith("Dear Sir or Madam"):
                 p = doc.add_paragraph()
                 p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                 p.paragraph_format.space_after = Pt(12)
                 run = p.add_run(clean_para)
                 self._set_font_and_size(run)
                 continue

            # Determine section and formatting
            is_intro_heading = clean_para.startswith("INTRODUCTION OF ") and clean_para.endswith(" QUALIFICATION")
            is_legal_authority_heading = clean_para == "LEGAL AUTHORITY"
            is_argument_heading = clean_para == "ARGUMENT"
            
            # Update section tracking
            if is_legal_authority_heading:
                # Legal authority will be inserted from template, so skip processing the rest
                in_legal_authority = False
                in_legal_criteria_list = False
                in_argument = False
            elif is_argument_heading:
                in_legal_authority = False
                in_legal_criteria_list = False
                in_argument = True

            # Create paragraph ONLY if we haven't handled it in special blocks below
            p = None
            
            # Handle "INTRODUCTION OF..." heading - centered, bold, underlined, black (matching sample)
            if is_intro_heading:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)  # Sample has no space_before
                p.paragraph_format.space_after = Pt(12)  # Add space after heading
                run = p.add_run(clean_para.replace("\n", " "))
                run.bold = True
                run.underline = True  # Underlined like other headings
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                self._set_font_and_size(run)
                continue
            
            # Handle Exhibit reference in Intro (e.g., "EXHIBIT 1 ---- Title")
            # Or the special marker "<<INSERT_EXHIBIT_1_FROM_DB>>"
            if clean_para.strip() == "<<INSERT_EXHIBIT_1_FROM_DB>>":
                if db and user_id:
                    # Fetch Exhibit 1
                    # We look for the exhibit with criteria_id="personal_info" or simply Exhibit 1
                    exhibit_repo = ExhibitRepository(db)
                    exhibit = exhibit_repo.get_latest_by_criteria(user_id, "personal_info")
                    if not exhibit:
                        exhibit = exhibit_repo.get_latest_by_number(user_id, 1)
                    
                    if exhibit:
                        # Reduce space_after of the PREVIOUS paragraph to close the gap
                        if doc.paragraphs:
                             doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

                        # Use the user-provided snippet structure
                        p_title = doc.add_paragraph()
                        # Use LEFT alignment instead of CENTER to match user request "remove spaces before"
                        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # Reduce space_before to 0 as requested ("also the line spacing before it")
                        p_title.paragraph_format.space_before = Pt(0)
                        p_title.paragraph_format.space_after = Pt(6)
                        
                        clean_title = exhibit.title.replace("EXHIBIT", "").strip() 
                        # Remove number if it's there (e.g. "1 ---- Title")
                        clean_title = re.sub(r'^\d+\s*[-]+\s*', '', clean_title).strip()
                        
                        full_title = f"EXHIBIT {exhibit.exhibit_number} ---- {clean_title}"
                        run_title = p_title.add_run(full_title)
                        run_title.bold = True
                        self._set_font_and_size(run_title)

                        items = exhibit_repo.get_items_for_exhibit_ordered(exhibit.id)
                        for item in items:
                            p = doc.add_paragraph()
                            p.paragraph_format.left_indent = Inches(0.5)
                            p.paragraph_format.space_after = Pt(0)
                            run_prefix = p.add_run(f"{item.item_suffix}. ")
                            run_prefix.bold = True
                            self._set_font_and_size(run_prefix)
                            
                            name = self._get_item_name(db, item)
                            run_desc = p.add_run(f"Copy of {name}")
                            self._set_font_and_size(run_desc)
                continue

            # Legacy text-based handling (fallback if marker not found but text present)
            if clean_para.startswith("EXHIBIT ") and "----" in clean_para:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(clean_para)
                run.bold = True
                self._set_font_and_size(run)
                continue

            # Handle list items in Intro (e.g., "a. Copy of...")
            if re.match(r'^[a-z]\.\s', clean_para):
                p = doc.add_paragraph(style='List Paragraph')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(0)
                # Indentation is handled by style, but we can enforce if needed
                p.paragraph_format.left_indent = Inches(0.5)
                
                # Split prefix ("a. ") and content ("Copy of...")
                match = re.match(r'^([a-z]\.\s)(.*)', clean_para)
                if match:
                    prefix, content = match.groups()
                    run_prefix = p.add_run(prefix)
                    run_prefix.bold = True
                    self._set_font_and_size(run_prefix)
                    
                    run_content = p.add_run(content)
                    self._set_font_and_size(run_content)
                else:
                    run = p.add_run(clean_para)
                    self._set_font_and_size(run)
                continue

            # Handle "LEGAL AUTHORITY" - insert exact template from sample
            if is_legal_authority_heading:
                # Insert the exact legal authority paragraphs from template
                for template_para in LEGAL_AUTHORITY_PARAGRAPHS:
                    p = doc.add_paragraph(template_para['text'], style=template_para['style'])
                    
                    # Set alignment
                    if template_para['alignment'] == 'JUSTIFY':
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    elif template_para['alignment'] == 'CENTER':
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif template_para['alignment'] == 'LEFT':
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # None alignment stays as default
                    
                    # Set indentation for list paragraphs
                    if template_para.get('left_indent'):
                        p.paragraph_format.left_indent = Inches(template_para['left_indent'])
                    
                    # Set spacing
                    p.paragraph_format.space_before = Pt(0)
                    # Add space after paragraphs (except for list items which should be tighter)
                    if template_para['style'] == 'List Paragraph':
                        p.paragraph_format.space_after = Pt(6)  # Less space for list items
                    else:
                        p.paragraph_format.space_after = Pt(12)  # More space for regular paragraphs
                    
                    # Handle formatting for heading (LEGAL AUTHORITY)
                    if template_para['text'] == 'LEGAL AUTHORITY':
                        # Format like INTRODUCTION and ARGUMENT: centered, bold, underlined, black
                        p.paragraph_format.space_before = Pt(12)  # Blank space before the heading
                        p.paragraph_format.space_after = Pt(12)  # Space after heading
                        for run in p.runs:
                            if template_para.get('bold'):
                                run.bold = True
                            if template_para.get('underline'):
                                run.underline = True
                            if template_para.get('color') == 'black':
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                            self._set_font_and_size(run)
                    else:
                        # Regular paragraphs
                        if template_para.get('bold') and p.runs:
                            p.runs[0].bold = True
                        # Set font for all runs
                        for run in p.runs:
                            self._set_font_and_size(run)
                
                # Mark that we're past legal authority section
                in_legal_authority = False
                continue
            
            # Handle "ARGUMENT" heading - Normal style, CENTER alignment, bold, underlined (matching sample)
            if is_argument_heading:
                # Reduce space_after of the PREVIOUS paragraph to close the gap
                if doc.paragraphs:
                     doc.paragraphs[-1].paragraph_format.space_after = Pt(12) # Reduced from potentially larger default, though 12 is standard

                p = doc.add_paragraph(clean_para.replace("\n", " "), style='Normal')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)  # Add space after heading
                # Make it bold, underlined, and black
                for run in p.runs:
                    run.bold = True
                    run.underline = True
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                    self._set_font_and_size(run)
                continue
            
            # Skip legal authority content from intro_text since we use the template
            # (The template is inserted when we see "LEGAL AUTHORITY" heading)
            if in_legal_authority:
                # Skip all content until we hit ARGUMENT
                continue
            
            # Handle argument section items - use List Paragraph style (matching sample)
            if in_argument and not is_argument_heading:
                # Check if this looks like a criteria item in the argument section
                # The intro sentence "X hereby submits..." should also be List Paragraph
                first_word = clean_para.split()[0] if clean_para.split() else ""
                if (clean_para.startswith("Evidence of") or 
                    clean_para.startswith("Documentation of") or 
                    clean_para.startswith("Published material") or
                    "hereby submits evidence" in clean_para.lower()):
                    # Use Normal paragraph style so Word does not collapse spacing
                    # between consecutive list-style lines.
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(12)  # Keep one paragraph gap between list lines
                    p.paragraph_format.left_indent = Pt(0)   # Remove the leading indent/space
                    p.paragraph_format.first_line_indent = Pt(0) # Ensure no first line indent
                    
                    # Add the text as a run to the existing paragraph
                    run = p.add_run(clean_para.replace("\n", " "))
                    self._set_font_and_size(run)
                    continue
            
            # Default: regular paragraph, justified
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Handle indentation for list items or quoted sections (like "(1)", "(A)", "(i)", "a.")
            if re.match(r'^(\([0-9A-Za-z]+\)|[a-z]\.)', clean_para):
                p.paragraph_format.left_indent = Inches(0.25)
            else:
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.first_line_indent = Pt(0)
            
            # Add spacing between paragraphs
            p.paragraph_format.space_after = Pt(12)  # Space after regular paragraphs
            
            # Replace newlines within paragraph with spaces for clean formatting
            formatted_text = clean_para.replace("\n", " ")
            
            # Add the text run
            run = p.add_run(formatted_text)
            self._set_font_and_size(run)

        return doc

    def generate_petition_conclusion(self, conclusion_text: str) -> Document:
        """
        Generates the final petition conclusion section.
        Simpler format than intro: one centered heading + justified paragraphs.
        """
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        if not conclusion_text:
            return doc

        if isinstance(conclusion_text, dict):
            conclusion_text = str(conclusion_text.get("text", conclusion_text.get("content", str(conclusion_text))))
        elif not isinstance(conclusion_text, str):
            conclusion_text = str(conclusion_text) if conclusion_text is not None else ""

        if not conclusion_text.strip():
            return doc

        paragraphs = re.split(r'\n\s*\n', conclusion_text.strip())
        for para_text in paragraphs:
            clean_para = (para_text or "").strip().replace("**", "").replace("*", "")
            if not clean_para:
                continue

            is_heading = clean_para.upper() == "CONCLUSION"
            p = doc.add_paragraph()
            if is_heading:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)  # Start conclusion section with top spacing
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run("CONCLUSION")
                run.bold = True
                self._set_font_and_size(run)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(clean_para.replace("\n", " "))
                self._set_font_and_size(run)

        return doc

    def generate_section(self, db: Session, criteria_id: str, exhibits: List[UserExhibit], conclusion_text: str = None, user_name: str = "The Petitioner", section_intro: str = None) -> Document:
        """
        Generates a generic petition section following the structural pattern:
        Title -> Section Intro -> Exhibit Evidence Cluster (Summary + Narrative + Sources) -> Conclusion
        """
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # Normalize criteria_id for lookup
        lookup_id = criteria_id.lower().replace("criterion_", "")
        lookup_id = {
            "press": "published_material",
            "publishedmaterial": "published_material",
            "published materials": "published_material",
            "published_materials": "published_material",
        }.get(lookup_id, lookup_id)
        info = CRITERIA_INFO.get(lookup_id, {
            "header": criteria_id.replace("_", " ").upper(),
            "reg_text": ""
        })

        # 1. Section Title (Matching Sample Style: "A. Evidence of [Name]'s receipt...")
        # Use the section letter stored in the database for consistency across multiple runs
        section_letter = exhibits[0].section_letter if (exhibits and exhibits[0].section_letter) else "A"
        
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Add a consistent space before every section title to separate it from the previous section
        title_p.paragraph_format.space_before = Pt(18)
        
        # Prepare regulatory text with name replacement
        raw_reg = info.get("reg_text") or info.get("header") or criteria_id
        # Remove the legal citation from the end for the title (e.g. 8 C.F.R. § 204.5(h)(3)(i))
        clean_reg = re.sub(r'\s*\d+\s+C\.F\.R\.\s+§\s+.*$', '', raw_reg).strip()
        
        # Replace "the Petitioner’s" or "the Petitioner" with the user's name
        # Use case-insensitive regex for replacement to be safe and handle both straight/smart quotes
        display_name = user_name if user_name.endswith('s') else f"{user_name}'s"
        
        # Handle "the Petitioner's" or "The Petitioner's"
        formatted_reg = re.sub(r'\bthe Petitioner[’\']s\b', display_name, clean_reg, flags=re.IGNORECASE)
        # Handle "the Petitioner" or "The Petitioner"
        formatted_reg = re.sub(r'\bthe Petitioner\b', user_name, formatted_reg, flags=re.IGNORECASE)
        
        title_run = title_p.add_run(f"{section_letter}. {formatted_reg}")
        title_run.bold = True
        self._set_font_and_size(title_run)
        # Reduce space_after to be more compact
        title_p.paragraph_format.space_after = Pt(6)

        # 2. Section Intro / Summary Paragraph (Optional)
        if section_intro:
            intro_p = doc.add_paragraph()
            intro_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            intro_run = intro_p.add_run(section_intro)
            self._set_font_and_size(intro_run)
            intro_p.paragraph_format.space_after = Pt(12)

        # 3. Interleaved Exhibit Clusters (Summary -> Items -> Narrative -> Sources)
        for i, exhibit in enumerate(exhibits):
            # A. Exhibit Header
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(18)
            clean_title = exhibit.title.strip().strip('“”"').strip()
            full_title = f"EXHIBIT {exhibit.exhibit_number} ---- {clean_title}"
            run_title = p_title.add_run(full_title)
            run_title.bold = True
            self._set_font_and_size(run_title)

            # B. Items List
            exhibit_repo = ExhibitRepository(db)
            items = exhibit_repo.get_items_for_exhibit_ordered(exhibit.id)
            for item in items:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(0)
                run_prefix = p.add_run(f"{item.item_suffix}. ")
                run_prefix.bold = True
                
                name = self._get_item_name(db, item)
                run_desc = p.add_run(f"Copy of {name}")
                self._set_font_and_size(run_desc)

            # C. Narrative Argument
            if exhibit.draft_content:
                # Add a small spacer before narrative
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                
                citation_repo = CitationRepository(db)
                citations = citation_repo.list_for_exhibit(exhibit.id)
                cite_map = {idx+1: c for idx, c in enumerate(citations)}
                
                # Split by one or more newlines that have only whitespace between them
                # This is more robust than just .split('\n\n')
                paragraphs = re.split(r'\n\s*\n', exhibit.draft_content.strip())
                for para_text in paragraphs:
                    if not para_text.strip(): continue
                    
                    # Cleanup and Normalization
                    # 1. Remove markdown bold/italic
                    clean_para = para_text.strip().replace("**", "").replace("*", "")
                    
                    # 2. Robustly remove "EXHIBIT TITLE: ..." labels
                    # We only strip it if the paragraph actually starts with that specific label
                    if clean_para.upper().startswith("EXHIBIT TITLE:"):
                        # Split by newline and remove only the first line
                        lines = clean_para.split('\n')
                        if len(lines) > 1:
                            clean_para = '\n'.join(lines[1:]).strip()
                        else:
                            # If it was ONLY the title line, make it empty so it's skipped
                            clean_para = ""
                    
                    if not clean_para: continue
                    
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # Internal Citation Cleanup
                    processed_text = clean_para
                    processed_text = re.sub(r'\[cite:\s*([^\]]+)\]', r'\1', processed_text)
                    processed_text = re.sub(r'\[(EXHIBIT\s+\d+\([a-z]\)[^\]]*)\]', r'\1', processed_text)
                    processed_text = re.sub(r'(?<!EXHIBIT\s)(\d+\([a-z]\))', r'EXHIBIT \1', processed_text)
                    
                    pattern = r'(\[\d+\])|(EXHIBIT\s+\d+\([a-z]\)(?:\s+(?:and|&|,\s*)\s*(?:EXHIBIT\s+)?\d+\([a-z]\))*)'
                    parts = re.split(pattern, processed_text)
                    
                    for part in parts:
                        if not part: continue
                        if part.startswith('[') and part.endswith(']') and part[1:-1].isdigit():
                            idx = int(part[1:-1])
                            if idx in cite_map:
                                run = p.add_run(str(idx))
                                run.font.superscript = True
                                run.bold = True
                                run.font.size = Pt(9)
                        elif "EXHIBIT" in part:
                            ref_run = p.add_run(part)
                            ref_run.bold = True
                            self._set_font_and_size(ref_run)
                        else:
                            run = p.add_run(part)
                            self._set_font_and_size(run)

                # D. Footnote Sources for THIS exhibit
                if citations:
                    line_p = doc.add_paragraph()
                    line_run = line_p.add_run("________________________")
                    line_run.font.size = Pt(10)
                    line_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    line_p.paragraph_format.space_before = Pt(6)
                    
                    h_p = doc.add_paragraph()
                    h_run = h_p.add_run("Sources")
                    h_run.bold = True
                    h_run.font.size = Pt(9)
                    h_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    h_p.paragraph_format.space_after = Pt(2)

                    for idx, citation in enumerate(citations):
                        cp = doc.add_paragraph()
                        cp.paragraph_format.left_indent = Inches(0.2)
                        cp.paragraph_format.space_after = Pt(2)
                        num_run = cp.add_run(f"[{idx+1}] ")
                        num_run.bold = True
                        num_run.font.size = Pt(8)
                        num_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                        title_run = cp.add_run(f"{citation.title} ")
                        title_run.font.size = Pt(8)
                        title_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                        self._add_hyperlink(cp, f"({citation.url})", citation.url, color="0000EE", underline=True)

        # 4. Section Conclusion
        if conclusion_text:
            conc_p = doc.add_paragraph()
            conc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            conc_run = conc_p.add_run(conclusion_text)
            self._set_font_and_size(conc_run)
            conc_p.paragraph_format.space_before = Pt(12) # Controlled gap before conclusion
            conc_p.paragraph_format.space_after = Pt(0)

        return doc

    def combine_sections(self, section_docs: List[Document]) -> io.BytesIO:
        """
        Combines multiple criteria documents into one final petition correctly.
        Uses a robust merging strategy to avoid XML corruption and preserve relationships.
        """
        if not section_docs:
            return io.BytesIO()

        # Start with the first document
        final_doc = section_docs[0]
        
        # Append subsequent documents
        for i in range(1, len(section_docs)):
            sub_doc = section_docs[i]
            
            # Check if sub_doc has any content before adding a break and appending
            has_content = False
            for element in sub_doc.element.body:
                if element.tag.endswith('p') or element.tag.endswith('tbl'):
                    # Check if it's more than just an empty paragraph
                    if element.tag.endswith('p'):
                        if len(element.xpath('.//w:t')) > 0:
                            has_content = True
                            break
                    else:
                        has_content = True
                        break
            
            if not has_content:
                logger.info(f"Skipping empty section document at index {i}")
                continue

            # Copy elements from sub_doc to final_doc
            for element in sub_doc.element.body:
                # We skip the sectPr (section properties) to maintain a continuous flow
                # but keep paragraphs, tables, etc.
                if element.tag.endswith('sectPr'):
                    continue
                final_doc.element.body.append(element)
            
            # CRITICAL: Copy relationships (hyperlinks, images, etc.)
            # This prevents the "Unreadable Content" error in Word
            for rel_id, rel in sub_doc.part.rels.items():
                if "hyperlink" in rel.reltype:
                    # Re-map the hyperlink relationship in the final document
                    final_doc.part.relate_to(rel.target_ref, rel.reltype, is_external=True)

        # Set 1 inch margins on all sections of the final document
        for section in final_doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        target = io.BytesIO()
        final_doc.save(target)
        target.seek(0)
        return target

    def assemble_petition_section(self, db: Session, exhibits: List[UserExhibit], criteria_id: str, conclusion_text: str = None, user_name: str = "The Petitioner", section_intro: str = None) -> io.BytesIO:
        """Entry point for generating a criteria section."""
        doc = self.generate_section(db, criteria_id, exhibits, conclusion_text, user_name, section_intro)
        target = io.BytesIO()
        doc.save(target)
        target.seek(0)
        return target

document_builder = DocumentBuilder()
